from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import RGBColor
import sys

def format_date(raw_date_str):
    month_map = {
        'Jan': '01', 'Feb': '02', 'Mar': '03', 'Apr': '04', 'May': '05', 'Jun': '06',
        'Jul': '07', 'Aug': '08', 'Sep': '09', 'Oct': '10', 'Nov': '11', 'Dec': '12'
    }
    try:
        cleaned_str = raw_date_str.replace(',', '').strip()
        parts = cleaned_str.split()
        if len(parts) < 3:
            return None
        day_str, month_str, time_str = parts[0], parts[1], parts[2]
        if month_str not in month_map:
            return None
        hour, minute = map(int, time_str.split(':'))
        current_year = datetime.now().year
        original_dt = datetime(
            year=current_year,
            month=int(month_map[month_str]),
            day=int(day_str),
            hour=hour,
            minute=minute
        )
        adjusted_dt = original_dt - timedelta(hours=1)
        return f"{adjusted_dt.strftime('%d.%m.%Y')} о {adjusted_dt.strftime('%H:%M')}"
    except Exception as e:
        print(f"Помилка форматування дати '{raw_date_str}': {e}")
        return None

def add_hyperlink(paragraph, text, url):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.underline = True

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    r = run._r
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def scrape_tgstat():
    try:
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service)
    except Exception as e:
        print(f"Помилка при запуску Chrome/WebDriver: {e}")
        print("Переконайтеся, що Chrome встановлено.")
        sys.exit(1)

    driver.get("https://tgstat.com/ru/search")
    print("-" * 50)
    print("🚀 Вікно Chrome відкрито.")
    print("Налаштуйте фільтри у цьому вікні.")
    print("Дочекайтеся повного завантаження результатів.")
    print("-" * 50)
    input("Після завантаження результатів натисніть Enter...")

    page_source = driver.page_source
    driver.quit()

    soup = BeautifulSoup(page_source, 'html.parser')
    posts = soup.select('.posts-list.lm-list-container div[id^="post-"]')
    if not posts:
        print("Не знайдено жодних постів.")
        return

    print(f"Знайдено {len(posts)} постів. Починаємо обробку...")

    doc = Document()
    doc.add_heading('Результати парсингу TGStat', level=1)
    processed_count = 0

    for post in posts:
        try:
            date_el = post.select_one('.text-muted.m-0')
            if not date_el:
                continue
            formatted_date = format_date(date_el.text.strip())
            if not formatted_date:
                continue

            name_el = post.select_one('h5[class="m-0"] a')
            if not name_el:
                continue
            channel_name = name_el.text.strip()

            link_el = post.select_one('div[class="ml-auto"] a')
            if not link_el or not link_el.has_attr('href'):
                continue
            raw_href = link_el['href']
            telegram_link = raw_href.replace("/channel/@", "https://t.me/")
            if not telegram_link.startswith("https://t.me/"):
                continue

            paragraph = doc.add_paragraph()
            text_before = f'{formatted_date} на Telegram-каналі під назвою "{channel_name}" за посиланням: '
            paragraph.add_run(text_before)
            add_hyperlink(paragraph, telegram_link, telegram_link)

            processed_count += 1
        except Exception as e:
            print(f"Помилка при обробці поста: {e}")
            continue

    filename = "tgstat_results.docx"
    if processed_count > 0:
        doc.save(filename)
        print("=" * 50)
        print(f"✅ Збережено {processed_count} записів у файл: {filename}")
        print("=" * 50)
    else:
        print("=" * 50)
        print("⚠️ Не було оброблено жодного поста.")
        print("=" * 50)

if __name__ == "__main__":
    scrape_tgstat()
